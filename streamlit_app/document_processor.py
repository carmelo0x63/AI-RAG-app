import os
import logging
from typing import List, Dict, Any, Optional
import PyPDF2
import docx
from io import BytesIO
import tiktoken

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.encoding = tiktoken.get_encoding("cl100k_base")  # GPT-4 encoding
    
    def process_uploaded_file(self, uploaded_file) -> List[Dict[str, Any]]:
        """Process an uploaded file and return chunks."""
        try:
            file_content = uploaded_file.read()
            filename = uploaded_file.name
            file_type = self._get_file_type(filename)
            
            # Extract text based on file type
            if file_type == 'pdf':
                text = self._extract_pdf_text(file_content)
            elif file_type == 'docx':
                text = self._extract_docx_text(file_content)
            elif file_type == 'txt':
                text = file_content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not text.strip():
                raise ValueError("No text content found in the file")
            
            # Split into chunks
            chunks = self._split_text_into_chunks(text)
            
            # Create document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc = {
                    'content': chunk,
                    'metadata': {
                        'filename': filename,
                        'file_type': file_type,
                        'chunk_index': i,
                        'total_chunks': len(chunks)
                    }
                }
                documents.append(doc)
            
            logger.info(f"Processed {filename}: {len(documents)} chunks created")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.name}: {e}")
            raise
    
    def _get_file_type(self, filename: str) -> str:
        """Determine file type from filename."""
        extension = filename.lower().split('.')[-1]
        if extension == 'pdf':
            return 'pdf'
        elif extension in ['docx', 'doc']:
            return 'docx'
        elif extension == 'txt':
            return 'txt'
        else:
            raise ValueError(f"Unsupported file extension: {extension}")
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc_file = BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """Split text into chunks based on token count."""
        try:
            # Tokenize the text
            tokens = self.encoding.encode(text)
            
            chunks = []
            start = 0
            
            while start < len(tokens):
                # Calculate end position
                end = start + self.chunk_size
                
                # If this is not the last chunk, find a good breaking point
                if end < len(tokens):
                    # Look for sentence endings near the chunk boundary
                    chunk_tokens = tokens[start:end]
                    chunk_text = self.encoding.decode(chunk_tokens)
                    
                    # Try to break at sentence boundaries
                    sentences = chunk_text.split('.')
                    if len(sentences) > 1:
                        # Keep all but the last incomplete sentence
                        chunk_text = '.'.join(sentences[:-1]) + '.'
                        chunk_tokens = self.encoding.encode(chunk_text)
                        end = start + len(chunk_tokens)
                
                # Extract the chunk
                chunk_tokens = tokens[start:end]
                chunk_text = self.encoding.decode(chunk_tokens)
                
                if chunk_text.strip():
                    chunks.append(chunk_text.strip())
                
                # Move to next chunk with overlap
                start = end - self.chunk_overlap
                
                # Prevent infinite loop
                if start >= end:
                    start = end
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error splitting text into chunks: {e}")
            # Fallback to simple character-based chunking
            return self._simple_text_split(text)
    
    def _simple_text_split(self, text: str) -> List[str]:
        """Simple fallback text splitting based on character count."""
        char_chunk_size = self.chunk_size * 4  # Rough estimate: 4 chars per token
        char_overlap = self.chunk_overlap * 4
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + char_chunk_size
            chunk = text[start:end]
            
            if chunk.strip():
                chunks.append(chunk.strip())
            
            start = end - char_overlap
            
            # Prevent infinite loop
            if start >= end:
                start = end
        
        return chunks
    
    def get_text_stats(self, text: str) -> Dict[str, Any]:
        """Get statistics about the text."""
        try:
            tokens = self.encoding.encode(text)
            return {
                'character_count': len(text),
                'word_count': len(text.split()),
                'token_count': len(tokens),
                'estimated_chunks': max(1, len(tokens) // self.chunk_size)
            }
        except Exception as e:
            logger.error(f"Error getting text stats: {e}")
            return {
                'character_count': len(text),
                'word_count': len(text.split()),
                'token_count': 0,
                'estimated_chunks': 1
            }